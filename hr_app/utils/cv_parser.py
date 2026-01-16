# hr_app/utils/cv_parser.py
import os
import re
from typing import List, Dict, Tuple

# Try multiple PDF libraries
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class CVParser:
    """
    Utility class to parse CV files and extract text.
    Supports PDF and DOCX formats with multiple fallback methods.
    """
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """
        Extract text from PDF using multiple methods for best results.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text as string
        """
        text = ""
        
        # Method 1: Try pdfplumber (best for most PDFs)
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    print("✓ Used pdfplumber for PDF extraction")
                    return text
            except Exception as e:
                print(f"pdfplumber failed: {e}")
        
        # Method 2: Try PyPDF2 (fallback)
        if PYPDF2_AVAILABLE:
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    print("✓ Used PyPDF2 for PDF extraction")
                    return text
            except Exception as e:
                print(f"PyPDF2 failed: {e}")
        
        # Method 3: Try pdfminer (most robust but slower)
        if PDFMINER_AVAILABLE:
            try:
                text = pdfminer_extract(file_path)
                if text.strip():
                    print("✓ Used pdfminer for PDF extraction")
                    return text
            except Exception as e:
                print(f"pdfminer failed: {e}")
        
        # If all methods fail, try OCR suggestion
        if not text.strip():
            print("⚠ No text could be extracted. The PDF might be scanned.")
            print("   Consider using OCR: pip install pytesseract pillow")
        
        return text
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """
        Extract text from DOCX file using python-docx.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text as string
        """
        text = ""
        if DOCX_AVAILABLE:
            try:
                doc = docx.Document(file_path)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\n"
            except Exception as e:
                print(f"Error reading DOCX: {e}")
        
        return text
    
    @staticmethod
    def extract_text(file_path: str, file_extension: str) -> str:
        """
        Main method to extract text based on file extension.
        
        Args:
            file_path: Path to the file
            file_extension: File extension (pdf, docx, etc.)
            
        Returns:
            Extracted text as string
        """
        file_extension = file_extension.lower()
        
        if file_extension == '.pdf':
            return CVParser.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return CVParser.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    @staticmethod
    def clean_extracted_text(text: str) -> str:
        """
        Clean and normalize extracted text for better analysis.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important ones
        text = re.sub(r'[^\w\s.,:;!?@()\-/]', ' ', text)
        
        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Remove multiple consecutive newlines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()


class CVAnalyzer:
    """
    Analyzes extracted CV text using rule-based logic.
    Updated with better parsing for accurate information extraction.
    """
    
    # Enhanced skill keywords with more variations
    SKILL_KEYWORDS = {
        'python': ['python', 'python3', 'python 3', 'python programming', 'django', 'flask', 'fastapi', 'pandas', 'numpy'],
        'java': ['java', 'java programming', 'spring', 'spring boot', 'spring framework', 'hibernate', 'j2ee'],
        'javascript': ['javascript', 'js', 'node.js', 'nodejs', 'react', 'angular', 'vue', 'typescript', 'express.js'],
        'sql': ['sql', 'mysql', 'postgresql', 'postgres', 'oracle', 'sql server', 'database', 'mongodb'],
        'excel': ['excel', 'microsoft excel', 'spreadsheet', 'vlookup', 'pivot table'],
        'ml': ['machine learning', 'ml', 'ai', 'artificial intelligence', 'deep learning', 'neural network', 'tensorflow', 'pytorch'],
        'data_analysis': ['data analysis', 'data analytics', 'power bi', 'tableau', 'data visualization', 'business intelligence'],
        'web_development': ['html', 'css', 'bootstrap', 'web development', 'frontend', 'backend', 'full stack'],
        'cloud': ['aws', 'amazon web services', 'azure', 'google cloud', 'gcp', 'cloud computing', 'docker', 'kubernetes'],
        'testing': ['testing', 'selenium', 'junit', 'testng', 'automation testing', 'manual testing'],
        'devops': ['devops', 'ci/cd', 'jenkins', 'git', 'github', 'gitlab'],
    }
    
    # Education keywords with degrees
    EDUCATION_KEYWORDS = [
        'bachelor', "bachelor's", 'b.tech', 'b.e.', 'b.sc', 'b.com', 'ba', 'b.a.',
        'master', "master's", 'm.tech', 'm.sc', 'm.com', 'ma', 'm.a', 'mba', 
        'phd', 'doctorate', 'ph.d',
        'diploma', 'degree', 'certificate',
        'university', 'college', 'institute', 'school',
        'computer science', 'engineering', 'information technology'
    ]
    
    # Job title patterns
    JOB_TITLE_PATTERNS = [
        r'(?:senior|junior|lead|principal)?\s*(?:software|web|frontend|backend|full.?stack)\s*(?:developer|engineer)',
        r'(?:data\s*(?:scientist|analyst|engineer))',
        r'(?:ml|machine\s*learning)\s*(?:engineer|specialist)',
        r'(?:devops|cloud)\s*(?:engineer|architect)',
        r'(?:qa|test)\s*(?:engineer|analyst)',
    ]
    
    # Enhanced experience patterns
    EXPERIENCE_PATTERNS = [
        r'(\d+)\s*(?:years?|yrs?)\s*(?:\+)?\s*experience',  # "5 years experience"
        r'experience.*?(\d+)\s*(?:years?|yrs?)',  # "experience of 5 years"
        r'(\d+)\+?\s*years?',  # "5+ years"
        r'(\d+)\s*-\s*(\d+)\s*years?',  # "3-5 years"
        r'(\d+)\s*months?',  # "24 months" -> convert to years
    ]
    
    @staticmethod
    def extract_skills(text: str) -> List[str]:
        """
        Extract skills from CV text using keyword matching.
        More sophisticated approach.
        
        Args:
            text: Extracted CV text (lowercase recommended)
            
        Returns:
            List of found skills
        """
        text_lower = text.lower()
        found_skills = []
        
        # Clean text - remove common words that might cause false positives
        # (e.g., "python" in "I have experience with Python programming")
        
        # Check for each skill and its variations
        for skill, variations in CVAnalyzer.SKILL_KEYWORDS.items():
            for variation in variations:
                # Use word boundaries to avoid partial matches
                pattern = r'\b' + re.escape(variation) + r'\b'
                if re.search(pattern, text_lower, re.IGNORECASE):
                    if skill not in found_skills:
                        found_skills.append(skill)
                    break
        
        # Additional check for skills mentioned in context
        skill_contexts = [
            (r'skills?[\s:]*(.*?)(?:\n\n|\n\w+:)', text_lower),  # Skills section
            (r'technical\s*skills?[\s:]*(.*?)(?:\n\n|\n\w+:)', text_lower),  # Technical skills
            (r'expertise[\s:]*(.*?)(?:\n\n|\n\w+:)', text_lower),  # Expertise section
        ]
        
        for pattern, search_text in skill_contexts:
            match = re.search(pattern, search_text, re.DOTALL | re.IGNORECASE)
            if match:
                skills_text = match.group(1)
                # Check for each skill in the skills section
                for skill, variations in CVAnalyzer.SKILL_KEYWORDS.items():
                    for variation in variations:
                        if variation in skills_text.lower():
                            if skill not in found_skills:
                                found_skills.append(skill)
                            break
        
        return found_skills
    
    @staticmethod
    def extract_experience(text: str) -> float:
        """
        Extract years of experience from CV text.
        More accurate implementation.
        
        Args:
            text: Extracted CV text
            
        Returns:
            Years of experience (float)
        """
        text_lower = text.lower()
        years_found = []
        
        # Pattern 1: Look for experience patterns
        for pattern in CVAnalyzer.EXPERIENCE_PATTERNS:
            matches = re.findall(pattern, text_lower)
            for match in matches:
                if isinstance(match, tuple):  # For patterns like "3-5 years"
                    try:
                        # Take average of range
                        min_years = float(match[0])
                        max_years = float(match[1]) if len(match) > 1 else min_years
                        years = (min_years + max_years) / 2
                        years_found.append(years)
                    except ValueError:
                        continue
                else:
                    try:
                        years = float(match)
                        if "months" in pattern:
                            years = years / 12  # Convert months to years
                        years_found.append(years)
                    except ValueError:
                        continue
        
        # Pattern 2: Look for work history dates
        date_patterns = [
            r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}\s*(?:-|to)\s*(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+\d{4}',
            r'\d{1,2}/\d{4}\s*(?:-|to)\s*\d{1,2}/\d{4}',
            r'\d{4}\s*(?:-|to)\s*\d{4}',
        ]
        
        for pattern in date_patterns:
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            if len(matches) >= 2:
                # If multiple date ranges found, calculate total years
                try:
                    # Simple estimation: each job ~2 years if not specified
                    total_years = len(matches) * 2
                    years_found.append(total_years)
                except:
                    pass
        
        # Pattern 3: Look for experience in summary/objective
        summary_pattern = r'(?:summary|objective|profile)[\s:]*(.*?)(?:\n\n|\n\w+:)'
        match = re.search(summary_pattern, text_lower, re.DOTALL | re.IGNORECASE)
        if match:
            summary_text = match.group(1)
            for pattern in CVAnalyzer.EXPERIENCE_PATTERNS:
                sub_matches = re.findall(pattern, summary_text)
                for match_val in sub_matches:
                    try:
                        years = float(match_val[0] if isinstance(match_val, tuple) else match_val)
                        years_found.append(years)
                    except:
                        continue
        
        # Return the maximum years found, or 0 if none
        return max(years_found) if years_found else 0.0
    
    @staticmethod
    def extract_education(text: str) -> str:
        """
        Extract education information from CV text.
        
        Args:
            text: Extracted CV text
            
        Returns:
            Education details as string
        """
        education_lines = []
        lines = text.split('\n')
        
        # First, try to find education section
        education_section = ""
        for i, line in enumerate(lines):
            line_lower = line.lower()
            if any(keyword in line_lower for keyword in ['education', 'academic', 'qualification']):
                # Take next few lines as education
                for j in range(i, min(i + 10, len(lines))):
                    education_section += lines[j] + "\n"
                break
        
        # If education section found, extract from it
        if education_section:
            section_lines = education_section.split('\n')
            for line in section_lines:
                line_lower = line.lower()
                # Check if line contains education keywords
                if any(keyword in line_lower for keyword in CVAnalyzer.EDUCATION_KEYWORDS):
                    if line.strip() and len(line.strip()) > 10:  # Avoid very short lines
                        education_lines.append(line.strip())
        else:
            # Fallback: scan entire text
            for line in lines:
                line_lower = line.lower()
                if any(keyword in line_lower for keyword in CVAnalyzer.EDUCATION_KEYWORDS):
                    if line.strip() and len(line.strip()) > 10:
                        education_lines.append(line.strip())
        
        # Remove duplicates while preserving order
        seen = set()
        unique_lines = []
        for line in education_lines:
            if line not in seen:
                seen.add(line)
                unique_lines.append(line)
        
        return ' | '.join(unique_lines[:3])  # Return top 3 education lines
    
    @staticmethod
    def extract_name(text: str) -> str:
        """
        Extract candidate name from CV text (simple heuristic).
        
        Args:
            text: Extracted CV text
            
        Returns:
            Candidate name if found, else empty string
        """
        lines = text.strip().split('\n')
        if lines:
            # First non-empty line is often the name
            for line in lines[:5]:  # Check first 5 lines
                line = line.strip()
                if line and len(line) < 50:  # Reasonable name length
                    # Check if it looks like a name (no numbers, not all caps)
                    if not re.search(r'\d', line) and not line.isupper():
                        # Remove common non-name prefixes
                        if not line.lower().startswith(('curriculum vitae', 'resume', 'cv', 'phone', 'email')):
                            return line
        return ""
    
    @staticmethod
    def extract_email(text: str) -> str:
        """
        Extract email from CV text.
        
        Args:
            text: Extracted CV text
            
        Returns:
            Email address if found, else empty string
        """
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        match = re.search(email_pattern, text)
        return match.group(0) if match else ""
    
    @staticmethod
    def extract_phone(text: str) -> str:
        """
        Extract phone number from CV text.
        
        Args:
            text: Extracted CV text
            
        Returns:
            Phone number if found, else empty string
        """
        phone_patterns = [
            r'\+\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',  # International
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # US/Canada style
            r'\d{10}',  # 10 digit number
        ]
        
        for pattern in phone_patterns:
            match = re.search(pattern, text)
            if match:
                return match.group(0)
        return ""
    
    @staticmethod
    def categorize_candidate(skills: List[str], experience_years: float) -> str:
        """
        Categorize candidate based on skills and experience.
        This implements the rule-based categorization logic.
        
        Args:
            skills: List of candidate's skills
            experience_years: Years of experience
            
        Returns:
            Category name as string
        """
        # Define category rules with priority
        category_rules = [
            {
                'name': 'Senior Python Developer',
                'condition': lambda s, e: 'python' in s and e >= 5 and ('django' in s or 'flask' in s)
            },
            {
                'name': 'Python Developer',
                'condition': lambda s, e: 'python' in s and ('django' in s or 'flask' in s)
            },
            {
                'name': 'Data Scientist',
                'condition': lambda s, e: 'ml' in s and ('python' in s or 'r' in s) and e >= 2
            },
            {
                'name': 'Data Analyst',
                'condition': lambda s, e: ('excel' in s and 'sql' in s) or 'data_analysis' in s
            },
            {
                'name': 'Senior Java Developer',
                'condition': lambda s, e: 'java' in s and e >= 5
            },
            {
                'name': 'Java Developer',
                'condition': lambda s, e: 'java' in s and e >= 1
            },
            {
                'name': 'Web Developer',
                'condition': lambda s, e: 'web_development' in s or ('javascript' in s and 'html' in s)
            },
            {
                'name': 'Full Stack Developer',
                'condition': lambda s, e: ('python' in s or 'java' in s or 'javascript' in s) and 
                                         ('web_development' in s or 'html' in s) and e >= 2
            },
            {
                'name': 'ML Engineer',
                'condition': lambda s, e: 'ml' in s and e >= 2
            },
            {
                'name': 'Cloud Engineer',
                'condition': lambda s, e: 'cloud' in s and e >= 2
            },
            {
                'name': 'DevOps Engineer',
                'condition': lambda s, e: 'devops' in s and e >= 2
            },
        ]
        
        # Check each category rule
        for rule in category_rules:
            if rule['condition'](skills, experience_years):
                return rule['name']
        
        # Default categories based on experience
        if experience_years == 0:
            return 'Fresher'
        elif experience_years < 2:
            return 'Junior Developer'
        elif experience_years < 5:
            return 'Mid-Level Developer'
        else:
            return 'Senior Professional'
        

# Add this to CVParser class

@staticmethod
def extract_text_with_ocr(file_path: str) -> str:
    """
    Extract text from scanned PDF using OCR.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text as string
    """
    try:
        import pytesseract
        from pdf2image import convert_from_path
        import tempfile
        
        text = ""
        
        # Convert PDF to images
        images = convert_from_path(file_path)
        
        # Extract text from each image
        for i, image in enumerate(images):
            page_text = pytesseract.image_to_string(image)
            text += f"--- Page {i+1} ---\n"
            text += page_text + "\n\n"
        
        return text
    except ImportError:
        print("OCR libraries not installed. Install with: pip install pytesseract pillow pdf2image")
        return ""
    except Exception as e:
        print(f"OCR failed: {e}")
        return ""